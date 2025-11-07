"""
Text Extraction and NER Module

This module handles text extraction from various file types (PDF, Word, JSON)
and performs Named Entity Recognition using AWS services.
"""

import boto3
import json
import logging
import tempfile
import os
from typing import Dict, List, Any, Optional, Union
from botocore.exceptions import ClientError
from fastapi import UploadFile
import docx
import PyPDF2
from io import BytesIO

logger = logging.getLogger(__name__)


class AWSTextExtractor:
    """
    Text extraction and NER service using AWS Textract and Comprehend Medical.
    """
    
    def __init__(
        self,
        aws_access_key_id: Optional[str] = None,
        aws_secret_access_key: Optional[str] = None,
        region_name: str = "us-east-1",
        aws_role_arn: Optional[str] = None
    ):
        """
        Initialize AWS Text Extractor.
        
        Args:
            aws_access_key_id: AWS access key ID
            aws_secret_access_key: AWS secret access key
            region_name: AWS region name
            aws_role_arn: AWS role ARN for role-based access
        """
        self.region_name = region_name
        self.aws_role_arn = aws_role_arn
        
        # Initialize AWS clients
        self._init_aws_clients(aws_access_key_id, aws_secret_access_key)
    
    def _init_aws_clients(
        self,
        aws_access_key_id: Optional[str],
        aws_secret_access_key: Optional[str]
    ):
        """Initialize AWS service clients."""
        try:
            session_kwargs = {"region_name": self.region_name}
            
            if aws_access_key_id and aws_secret_access_key:
                session_kwargs.update({
                    "aws_access_key_id": aws_access_key_id,
                    "aws_secret_access_key": aws_secret_access_key
                })
            
            # Create session
            session = boto3.Session(**session_kwargs)
            
            # Assume role if specified
            if self.aws_role_arn:
                sts_client = session.client('sts')
                assumed_role = sts_client.assume_role(
                    RoleArn=self.aws_role_arn,
                    RoleSessionName='text-extraction-session'
                )
                
                credentials = assumed_role['Credentials']
                session = boto3.Session(
                    aws_access_key_id=credentials['AccessKeyId'],
                    aws_secret_access_key=credentials['SecretAccessKey'],
                    aws_session_token=credentials['SessionToken'],
                    region_name=self.region_name
                )
            
            # Initialize service clients
            self.textract_client = session.client('textract')
            self.comprehend_medical_client = session.client('comprehendmedical')
            
            logger.info("AWS clients initialized successfully")
            
        except Exception as e:
            logger.error(f"Failed to initialize AWS clients: {str(e)}")
            raise
    
    async def extract_and_analyze(
        self,
        file: Union[UploadFile, bytes, str],
        file_type: str,
        include_ner: bool = True
    ) -> Dict[str, Any]:
        """
        Extract text and perform NER analysis on the input file.
        
        Args:
            file: Input file (UploadFile, bytes, or file path)
            file_type: Type of file (pdf, docx, json, txt)
            include_ner: Whether to include NER analysis
            
        Returns:
            Dictionary containing extracted text and NER results
        """
        try:
            # Extract text based on file type
            extracted_text = await self._extract_text(file, file_type)
            
            result = {
                "extracted_text": extracted_text,
                "file_type": file_type,
                "text_length": len(extracted_text),
                "extraction_method": self._get_extraction_method(file_type)
            }
            
            # Perform NER if requested
            if include_ner and extracted_text:
                ner_results = await self._perform_ner(extracted_text)
                result["ner_analysis"] = ner_results
            
            return result
            
        except Exception as e:
            logger.error(f"Error in extract_and_analyze: {str(e)}")
            raise
    
    async def _extract_text(self, file: Union[UploadFile, bytes, str], file_type: str) -> str:
        """Extract text from file based on type."""
        file_type = file_type.lower()
        
        if file_type == "pdf":
            return await self._extract_from_pdf(file)
        elif file_type in ["docx", "doc"]:
            return await self._extract_from_word(file)
        elif file_type == "json":
            return await self._extract_from_json(file)
        elif file_type == "txt":
            return await self._extract_from_text(file)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def _extract_from_pdf(self, file: Union[UploadFile, bytes, str]) -> str:
        """Extract text from PDF using AWS Textract."""
        try:
            # Get file bytes
            if isinstance(file, UploadFile):
                file_bytes = await file.read()
            elif isinstance(file, bytes):
                file_bytes = file
            else:
                with open(file, 'rb') as f:
                    file_bytes = f.read()
            
            # Use Textract for PDF text extraction
            response = self.textract_client.detect_document_text(
                Document={'Bytes': file_bytes}
            )
            
            # Extract text from Textract response
            extracted_text = ""
            for block in response.get('Blocks', []):
                if block['BlockType'] == 'LINE':
                    extracted_text += block.get('Text', '') + "\n"
            
            return extracted_text.strip()
            
        except ClientError as e:
            logger.error(f"AWS Textract error: {str(e)}")
            # Fallback to PyPDF2 if Textract fails
            return await self._extract_pdf_fallback(file)
        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            return await self._extract_pdf_fallback(file)
    
    async def _extract_pdf_fallback(self, file: Union[UploadFile, bytes, str]) -> str:
        """Fallback PDF extraction using PyPDF2."""
        try:
            if isinstance(file, UploadFile):
                file_bytes = await file.read()
                pdf_file = BytesIO(file_bytes)
            elif isinstance(file, bytes):
                pdf_file = BytesIO(file)
            else:
                pdf_file = open(file, 'rb')
            
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            if not isinstance(file, (UploadFile, bytes)):
                pdf_file.close()
                
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF fallback extraction error: {str(e)}")
            return ""
    
    async def _extract_from_word(self, file: Union[UploadFile, bytes, str]) -> str:
        """Extract text from Word document."""
        try:
            if isinstance(file, UploadFile):
                # Save to temporary file for docx processing
                with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                    content = await file.read()
                    tmp_file.write(content)
                    tmp_file_path = tmp_file.name
                
                doc = docx.Document(tmp_file_path)
                os.unlink(tmp_file_path)  # Clean up temp file
                
            elif isinstance(file, bytes):
                doc = docx.Document(BytesIO(file))
            else:
                doc = docx.Document(file)
            
            # Extract text from all paragraphs
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Word document extraction error: {str(e)}")
            return ""
    
    async def _extract_from_json(self, file: Union[UploadFile, bytes, str]) -> str:
        """Extract text from JSON file."""
        try:
            if isinstance(file, UploadFile):
                content = await file.read()
                json_data = json.loads(content.decode('utf-8'))
            elif isinstance(file, bytes):
                json_data = json.loads(file.decode('utf-8'))
            else:
                with open(file, 'r', encoding='utf-8') as f:
                    json_data = json.load(f)
            
            # Convert JSON to readable text
            return json.dumps(json_data, indent=2, ensure_ascii=False)
            
        except Exception as e:
            logger.error(f"JSON extraction error: {str(e)}")
            return ""
    
    async def _extract_from_text(self, file: Union[UploadFile, bytes, str]) -> str:
        """Extract text from plain text file."""
        try:
            if isinstance(file, UploadFile):
                content = await file.read()
                return content.decode('utf-8')
            elif isinstance(file, bytes):
                return file.decode('utf-8')
            else:
                with open(file, 'r', encoding='utf-8') as f:
                    return f.read()
            
        except Exception as e:
            logger.error(f"Text file extraction error: {str(e)}")
            return ""
    
    async def _perform_ner(self, text: str) -> Dict[str, Any]:
        """Perform Named Entity Recognition using AWS Comprehend Medical."""
        try:
            # Limit text length for Comprehend Medical (max 20,000 bytes)
            max_length = 20000
            if len(text.encode('utf-8')) > max_length:
                text = text[:max_length]
                logger.warning("Text truncated for NER analysis due to length limit")
            
            # Use Comprehend Medical for medical NER
            entities_response = self.comprehend_medical_client.detect_entities_v2(
                Text=text
            )
            
            # Process entities
            entities = []
            for entity in entities_response.get('Entities', []):
                entities.append({
                    "text": entity.get('Text', ''),
                    "category": entity.get('Category', ''),
                    "type": entity.get('Type', ''),
                    "score": entity.get('Score', 0.0),
                    "begin_offset": entity.get('BeginOffset', 0),
                    "end_offset": entity.get('EndOffset', 0),
                    "attributes": entity.get('Attributes', [])
                })
            
            # Also detect PHI (Personal Health Information)
            phi_response = self.comprehend_medical_client.detect_phi(
                Text=text
            )
            
            phi_entities = []
            for entity in phi_response.get('Entities', []):
                phi_entities.append({
                    "text": entity.get('Text', ''),
                    "category": entity.get('Category', ''),
                    "type": entity.get('Type', ''),
                    "score": entity.get('Score', 0.0),
                    "begin_offset": entity.get('BeginOffset', 0),
                    "end_offset": entity.get('EndOffset', 0)
                })
            
            return {
                "medical_entities": entities,
                "phi_entities": phi_entities,
                "total_medical_entities": len(entities),
                "total_phi_entities": len(phi_entities),
                "categories": list(set([e["category"] for e in entities])),
                "types": list(set([e["type"] for e in entities]))
            }
            
        except ClientError as e:
            logger.error(f"AWS Comprehend Medical error: {str(e)}")
            return {
                "error": f"NER analysis failed: {str(e)}",
                "medical_entities": [],
                "phi_entities": [],
                "total_medical_entities": 0,
                "total_phi_entities": 0
            }
        except Exception as e:
            logger.error(f"NER analysis error: {str(e)}")
            return {
                "error": f"NER analysis failed: {str(e)}",
                "medical_entities": [],
                "phi_entities": [],
                "total_medical_entities": 0,
                "total_phi_entities": 0
            }
    
    def _get_extraction_method(self, file_type: str) -> str:
        """Get the extraction method used for the file type."""
        methods = {
            "pdf": "AWS Textract (with PyPDF2 fallback)",
            "docx": "python-docx",
            "doc": "python-docx",
            "json": "JSON parser",
            "txt": "Plain text reader"
        }
        return methods.get(file_type.lower(), "Unknown")